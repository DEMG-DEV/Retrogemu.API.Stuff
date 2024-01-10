from docx import Document
from docx.shared import Cm, Pt

doc = Document()

# Header
header = doc.sections[0].header
paragraph = header.paragraphs[0]
paragraph.add_run("Retrogemu docx file")

# Title
title = doc.add_heading("Title 1", level=1)
title_font = title.runs[0].font
title_font.size = Pt(18)

# Paragraph
paragraph_1 = doc.add_paragraph(
    "Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip")

# Bold text
paragraph_1.add_run(" ex ea commodo consequat.").bold = True
paragraph_1.add_run("Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.")

# Subtitle
subtitle_1 = doc.add_heading("Profile image", level=2)
doc.add_paragraph()
doc.add_picture('assets/img/profile_image.png', width=Cm(2), height=Cm(2))
paragraph_2 = doc.add_paragraph()
paragraph_2.add_run("@Retrogemu").bold = True

# Save file
doc.save("profile_example.docx")
